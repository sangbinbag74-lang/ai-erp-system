"""
Document Processor for AI ERP System

Processes various document types and extracts structured data
with AI-powered analysis and automatic content recognition.
"""

import os
import magic
import hashlib
import logging
from typing import Dict, List, Optional, Any, Tuple, Union
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
import asyncio
from pathlib import Path

# Document processing libraries
import PyPDF2
from PIL import Image
import pandas as pd
from docx import Document as DocxDocument
import openpyxl
import json
import csv
import io

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx" 
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    IMAGE = "image"
    TXT = "txt"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


@dataclass
class DocumentInfo:
    """Document metadata and processing information"""
    id: str
    filename: str
    file_type: DocumentType
    size: int
    mime_type: str
    hash: str
    uploaded_at: datetime
    processed_at: Optional[datetime] = None
    status: ProcessingStatus = ProcessingStatus.PENDING
    metadata: Dict[str, Any] = None
    error: Optional[str] = None


@dataclass
class ExtractedContent:
    """Extracted content from document"""
    text: str
    structured_data: Optional[Dict[str, Any]] = None
    images: List[bytes] = None
    tables: List[Dict[str, Any]] = None
    metadata: Dict[str, Any] = None


class DocumentProcessor:
    """
    AI-enhanced document processor for ERP files
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.max_file_size = self.config.get('max_file_size', 50 * 1024 * 1024)  # 50MB
        self.supported_types = {
            'application/pdf': DocumentType.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
            'text/csv': DocumentType.CSV,
            'application/json': DocumentType.JSON,
            'text/plain': DocumentType.TXT,
            'image/jpeg': DocumentType.IMAGE,
            'image/png': DocumentType.IMAGE,
            'image/tiff': DocumentType.IMAGE
        }
        
        # Initialize magic for file type detection
        self.magic = magic.Magic(mime=True)
    
    def _detect_file_type(self, file_data: bytes, filename: str) -> Tuple[str, DocumentType]:
        """Detect file type from content and filename"""
        
        try:
            # Use python-magic for MIME type detection
            mime_type = self.magic.from_buffer(file_data)
            
            # Map MIME type to document type
            doc_type = self.supported_types.get(mime_type, DocumentType.UNKNOWN)
            
            # Fallback to extension if MIME type is unknown
            if doc_type == DocumentType.UNKNOWN:
                ext = Path(filename).suffix.lower()
                extension_map = {
                    '.pdf': DocumentType.PDF,
                    '.docx': DocumentType.DOCX,
                    '.xlsx': DocumentType.XLSX,
                    '.csv': DocumentType.CSV,
                    '.json': DocumentType.JSON,
                    '.txt': DocumentType.TXT,
                    '.jpg': DocumentType.IMAGE,
                    '.jpeg': DocumentType.IMAGE,
                    '.png': DocumentType.IMAGE,
                    '.tiff': DocumentType.IMAGE
                }
                doc_type = extension_map.get(ext, DocumentType.UNKNOWN)
            
            return mime_type, doc_type
            
        except Exception as e:
            logger.warning(f"Error detecting file type: {str(e)}")
            return "application/octet-stream", DocumentType.UNKNOWN
    
    def _calculate_hash(self, file_data: bytes) -> str:
        """Calculate SHA256 hash of file"""
        return hashlib.sha256(file_data).hexdigest()
    
    async def process_document(
        self, 
        file_data: bytes, 
        filename: str,
        user_id: str,
        **kwargs
    ) -> Tuple[DocumentInfo, Optional[ExtractedContent]]:
        """
        Main document processing function
        """
        
        try:
            # Validate file size
            if len(file_data) > self.max_file_size:
                raise ValueError(f"File too large: {len(file_data)} bytes > {self.max_file_size}")
            
            # Detect file type
            mime_type, doc_type = self._detect_file_type(file_data, filename)
            
            # Create document info
            doc_id = f"{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{self._calculate_hash(file_data)[:8]}"
            
            doc_info = DocumentInfo(
                id=doc_id,
                filename=filename,
                file_type=doc_type,
                size=len(file_data),
                mime_type=mime_type,
                hash=self._calculate_hash(file_data),
                uploaded_at=datetime.now(),
                status=ProcessingStatus.PROCESSING,
                metadata={
                    "user_id": user_id,
                    "processing_options": kwargs
                }
            )
            
            logger.info(f"Processing document {doc_id}: {filename} ({doc_type.value})")
            
            # Extract content based on document type
            extracted_content = await self._extract_content(file_data, doc_type, doc_info)
            
            # Update status
            doc_info.status = ProcessingStatus.COMPLETED
            doc_info.processed_at = datetime.now()
            
            return doc_info, extracted_content
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            doc_info.status = ProcessingStatus.FAILED
            doc_info.error = str(e)
            return doc_info, None
    
    async def _extract_content(
        self, 
        file_data: bytes, 
        doc_type: DocumentType, 
        doc_info: DocumentInfo
    ) -> ExtractedContent:
        """Extract content based on document type"""
        
        try:
            if doc_type == DocumentType.PDF:
                return await self._process_pdf(file_data)
            elif doc_type == DocumentType.DOCX:
                return await self._process_docx(file_data)
            elif doc_type == DocumentType.XLSX:
                return await self._process_xlsx(file_data)
            elif doc_type == DocumentType.CSV:
                return await self._process_csv(file_data)
            elif doc_type == DocumentType.JSON:
                return await self._process_json(file_data)
            elif doc_type == DocumentType.TXT:
                return await self._process_txt(file_data)
            elif doc_type == DocumentType.IMAGE:
                return await self._process_image(file_data)
            else:
                logger.warning(f"Unsupported document type: {doc_type}")
                return ExtractedContent(
                    text="",
                    metadata={"error": "Unsupported document type"}
                )
                
        except Exception as e:
            logger.error(f"Error extracting content: {str(e)}")
            return ExtractedContent(
                text="",
                metadata={"error": str(e)}
            )
    
    async def _process_pdf(self, file_data: bytes) -> ExtractedContent:
        """Process PDF document"""
        
        try:
            pdf_file = io.BytesIO(file_data)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            metadata = {
                "page_count": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None
            }
            
            # Extract text from all pages
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            return ExtractedContent(
                text=text_content.strip(),
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    async def _process_docx(self, file_data: bytes) -> ExtractedContent:
        """Process DOCX document"""
        
        try:
            doc_file = io.BytesIO(file_data)
            doc = DocxDocument(doc_file)
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({"data": table_data})
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            return ExtractedContent(
                text=text_content.strip(),
                tables=tables if tables else None,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    async def _process_xlsx(self, file_data: bytes) -> ExtractedContent:
        """Process Excel spreadsheet"""
        
        try:
            excel_file = io.BytesIO(file_data)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            structured_data = {}
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to data
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):  # Skip empty rows
                        data.append(list(row))
                
                structured_data[sheet_name] = data
                
                # Create text representation
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                for row in data:
                    text_content += "\t".join(str(cell) if cell is not None else "" for cell in row) + "\n"
            
            metadata = {
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames
            }
            
            return ExtractedContent(
                text=text_content.strip(),
                structured_data=structured_data,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing XLSX: {str(e)}")
            raise
    
    async def _process_csv(self, file_data: bytes) -> ExtractedContent:
        """Process CSV file"""
        
        try:
            csv_content = file_data.decode('utf-8')
            csv_file = io.StringIO(csv_content)
            
            # Try to detect delimiter
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(csv_content[:1024]).delimiter
            
            csv_file.seek(0)
            reader = csv.reader(csv_file, delimiter=delimiter)
            
            data = []
            for row in reader:
                data.append(row)
            
            # Create text content
            text_content = ""
            for row in data:
                text_content += delimiter.join(row) + "\n"
            
            metadata = {
                "row_count": len(data),
                "column_count": len(data[0]) if data else 0,
                "delimiter": delimiter
            }
            
            return ExtractedContent(
                text=text_content.strip(),
                structured_data={"csv_data": data},
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing CSV: {str(e)}")
            raise
    
    async def _process_json(self, file_data: bytes) -> ExtractedContent:
        """Process JSON file"""
        
        try:
            json_content = file_data.decode('utf-8')
            data = json.loads(json_content)
            
            # Create readable text representation
            text_content = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                "json_type": type(data).__name__,
                "size": len(json_content)
            }
            
            if isinstance(data, list):
                metadata["item_count"] = len(data)
            elif isinstance(data, dict):
                metadata["key_count"] = len(data)
            
            return ExtractedContent(
                text=text_content,
                structured_data={"json_data": data},
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing JSON: {str(e)}")
            raise
    
    async def _process_txt(self, file_data: bytes) -> ExtractedContent:
        """Process text file"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    text_content = file_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise ValueError("Could not decode text file with supported encodings")
            
            metadata = {
                "character_count": len(text_content),
                "line_count": text_content.count('\n') + 1,
                "word_count": len(text_content.split())
            }
            
            return ExtractedContent(
                text=text_content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            raise
    
    async def _process_image(self, file_data: bytes) -> ExtractedContent:
        """Process image file"""
        
        try:
            image_file = io.BytesIO(file_data)
            image = Image.open(image_file)
            
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height
            }
            
            # Extract EXIF data if available
            if hasattr(image, '_getexif'):
                exif = image._getexif()
                if exif:
                    metadata["exif"] = dict(exif)
            
            return ExtractedContent(
                text=f"Image file: {image.format} {image.size[0]}x{image.size[1]}",
                images=[file_data],
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            raise
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        
        return {
            "supported_types": [dt.value for dt in DocumentType if dt != DocumentType.UNKNOWN],
            "max_file_size": self.max_file_size,
            "mime_types": list(self.supported_types.keys())
        }
    
    async def batch_process(
        self, 
        files: List[Tuple[bytes, str]], 
        user_id: str,
        max_concurrent: int = 5
    ) -> List[Tuple[DocumentInfo, Optional[ExtractedContent]]]:
        """Process multiple documents concurrently"""
        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_single(file_data, filename):
            async with semaphore:
                return await self.process_document(file_data, filename, user_id)
        
        tasks = [process_single(file_data, filename) for file_data, filename in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Error processing file {files[i][1]}: {str(result)}")
                error_doc = DocumentInfo(
                    id=f"error_{i}",
                    filename=files[i][1],
                    file_type=DocumentType.UNKNOWN,
                    size=len(files[i][0]),
                    mime_type="unknown",
                    hash="",
                    uploaded_at=datetime.now(),
                    status=ProcessingStatus.FAILED,
                    error=str(result)
                )
                processed_results.append((error_doc, None))
            else:
                processed_results.append(result)
        
        return processed_results


def create_document_processor(**config) -> DocumentProcessor:
    """Factory function to create document processor"""
    return DocumentProcessor(config)